# -*- coding: utf-8 -*-
"""
PDF -> DOCX, helt lokalt og selvforsynt.
Skannede dokumenter håndteres med innebygd Tesseract (ingen eksterne installasjoner).
"""

import os
import shutil
import subprocess
import sys
import tempfile

import fitz  # PyMuPDF
from pdf2docx import Converter

OCR_SPRAAK = "nor+eng"
OCR_DPI = 300


def _ressurs_sti(*deler) -> str:
    """Sti til medfølgende ressurser, både i utvikling og i PyInstaller-exe."""
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, *deler)


def finn_tesseract() -> str | None:
    """Innbakt Tesseract først, deretter system-PATH."""
    kandidater = [
        _ressurs_sti("tesseract", "tesseract.exe"),
        _ressurs_sti("tesseract", "tesseract"),
    ]
    for k in kandidater:
        if os.path.isfile(k):
            return k
    return shutil.which("tesseract")


def _tessdata_dir(tesseract_sti: str) -> str | None:
    """Finn tessdata-mappen som hører til den innbakte Tesseract."""
    lokal = os.path.join(os.path.dirname(tesseract_sti), "tessdata")
    return lokal if os.path.isdir(lokal) else None


def har_tekstlag(pdf_sti: str, min_tegn_per_side: int = 25) -> bool:
    """Digital PDF (ekte tekstlag) eller skannet (kun bilder)?"""
    try:
        doc = fitz.open(pdf_sti)
    except Exception:
        return False
    sider = min(len(doc), 10)
    tekst_sider = sum(
        1 for i in range(sider) if len(doc[i].get_text().strip()) >= min_tegn_per_side
    )
    doc.close()
    return sider > 0 and tekst_sider >= max(1, sider // 2)


def kjor_ocr(pdf_inn: str, pdf_ut: str, status=lambda s: None,
             fremdrift=lambda p: None, fra=10.0, til=60.0) -> None:
    """
    OCR uten Ghostscript: render hver side til bilde, la Tesseract lage
    søkbar PDF per side, og flett sidene sammen med PyMuPDF.
    """
    tess = finn_tesseract()
    if not tess:
        raise RuntimeError("Tesseract OCR mangler.")

    env = os.environ.copy()
    tessdata = _tessdata_dir(tess)
    if tessdata:
        env["TESSDATA_PREFIX"] = tessdata

    # Skjul konsollvindu på Windows
    flags = subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0

    kilde = fitz.open(pdf_inn)
    resultat = fitz.open()
    antall = len(kilde)

    with tempfile.TemporaryDirectory(prefix="pdf2word_") as tmp:
        for i in range(antall):
            status(f"Tekstgjenkjenning side {i + 1} av {antall} ...")
            pix = kilde[i].get_pixmap(dpi=OCR_DPI)
            bilde = os.path.join(tmp, f"s{i}.png")
            pix.save(bilde)
            del pix

            ut_base = os.path.join(tmp, f"s{i}")
            res = subprocess.run(
                [tess, bilde, ut_base, "-l", OCR_SPRAAK,
                 "--dpi", str(OCR_DPI), "pdf"],
                capture_output=True, text=True, env=env, creationflags=flags,
            )
            side_pdf = ut_base + ".pdf"
            if res.returncode != 0 or not os.path.isfile(side_pdf):
                raise RuntimeError(f"OCR feilet på side {i + 1}:\n{res.stderr[-400:]}")

            side_doc = fitz.open(side_pdf)
            resultat.insert_pdf(side_doc)
            side_doc.close()
            os.remove(bilde)
            fremdrift(fra + (i + 1) / antall * (til - fra))

        kilde.close()
        resultat.save(pdf_ut)
        resultat.close()


def bygg_docx_fra_ocr(ocr_pdf: str, docx_sti: str) -> None:
    """
    Bygger DOCX fra OCR-tekstlaget. pdf2docx ignorerer usynlig OCR-tekst,
    så vi rekonstruerer avsnitt og overskrifter selv fra tekstblokkene.
    """
    from docx import Document
    from docx.shared import Pt

    doc = fitz.open(ocr_pdf)
    word = Document()

    # Finn typisk brødtekststørrelse for å kunne skille ut overskrifter
    storrelser = []
    for side in doc:
        for blokk in side.get_text("dict")["blocks"]:
            for linje in blokk.get("lines", []):
                for span in linje.get("spans", []):
                    if span["text"].strip():
                        storrelser.append(span["size"])
    if not storrelser:
        doc.close()
        raise RuntimeError("Fant ingen tekst etter OCR. Skannet kan være for utydelig.")
    storrelser.sort()
    median = storrelser[len(storrelser) // 2]

    forste_side = True
    for side in doc:
        if not forste_side:
            word.add_page_break()
        forste_side = False

        for blokk in side.get_text("dict")["blocks"]:
            linjer = blokk.get("lines", [])
            if not linjer:
                continue
            tekst_linjer, maks_str = [], 0.0
            for linje in linjer:
                t = " ".join(s["text"] for s in linje.get("spans", [])).strip()
                if t:
                    tekst_linjer.append(t)
                    maks_str = max(maks_str, max(s["size"] for s in linje["spans"]))
            if not tekst_linjer:
                continue
            tekst = " ".join(" ".join(tekst_linjer).split())

            if maks_str >= median * 1.5 and len(tekst) < 120:
                word.add_heading(tekst, level=1)
            elif maks_str >= median * 1.2 and len(tekst) < 120:
                word.add_heading(tekst, level=2)
            else:
                avsnitt = word.add_paragraph(tekst)
                for run in avsnitt.runs:
                    run.font.size = Pt(11)

    doc.close()
    word.save(docx_sti)


def _normaliser(tekst: str) -> str:
    """Normaliserer tekst for sammenligning: sidetall o.l. blir like."""
    import re
    t = " ".join(tekst.split()).lower()
    return re.sub(r"\d+", "#", t)


def fjern_topp_bunn(pdf_inn: str, pdf_ut: str, sone: float = 0.10) -> int:
    """
    Fjerner topp- og bunntekst: tekstblokker i øvre/nedre sone som gjentas
    på tvers av sider, samt rene sidetall. Returnerer antall fjernede blokker.
    """
    import re
    doc = fitz.open(pdf_inn)
    antall_sider = len(doc)

    # Tell forekomster av normalisert tekst i topp-/bunnsonene
    forekomster: dict[str, int] = {}
    kandidater = []  # (side_idx, bbox, nokkel, tekst)
    for i, side in enumerate(doc):
        h = side.rect.height
        for blokk in side.get_text("blocks"):
            x0, y0, x1, y1, tekst = blokk[0], blokk[1], blokk[2], blokk[3], blokk[4]
            tekst = tekst.strip()
            if not tekst:
                continue
            i_topp = y1 <= h * 0.08          # hele blokken i øverste 8 %
            i_bunn = y0 >= h * 0.92          # hele blokken i nederste 8 %
            liten = (y1 - y0) <= 40 and len(tekst) <= 150  # aldri overskrifter/avsnitt
            if (i_topp or i_bunn) and liten:
                nokkel = _normaliser(tekst)
                forekomster[nokkel] = forekomster.get(nokkel, 0) + 1
                kandidater.append((i, fitz.Rect(x0, y0, x1, y1), nokkel, tekst))

    terskel = max(2, antall_sider // 2)
    sidetall_monster = re.compile(r"^(side\s*)?\d+(\s*(av|/)\s*\d+)?$", re.IGNORECASE)

    fjernet = 0
    per_side: dict[int, list] = {}
    for side_idx, bbox, nokkel, tekst in kandidater:
        gjentatt = antall_sider >= 3 and forekomster[nokkel] >= terskel
        sidetall = bool(sidetall_monster.match(" ".join(tekst.split())))
        if gjentatt or sidetall:
            per_side.setdefault(side_idx, []).append(bbox)
            fjernet += 1

    for side_idx, bokser in per_side.items():
        side = doc[side_idx]
        for bbox in bokser:
            side.add_redact_annot(bbox)
        side.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    doc.save(pdf_ut, garbage=3)
    doc.close()
    return fjernet


def ledig_filnavn(sti: str) -> str:
    """fil.docx -> fil (2).docx hvis den finnes fra før."""
    if not os.path.exists(sti):
        return sti
    base, ext = os.path.splitext(sti)
    n = 2
    while os.path.exists(f"{base} ({n}){ext}"):
        n += 1
    return f"{base} ({n}){ext}"


def konverter(pdf_sti: str, docx_sti: str | None = None,
              status=lambda s: None, fremdrift=lambda p: None,
              fjern_marger: bool = True) -> dict:
    """Hovedløp. Returnerer {'utfil': ..., 'ocr_brukt': bool, 'fjernet_blokker': int}."""
    if not os.path.isfile(pdf_sti):
        raise FileNotFoundError(f"Finner ikke filen: {pdf_sti}")
    if not pdf_sti.lower().endswith(".pdf"):
        raise ValueError("Filen må være en PDF.")

    if docx_sti is None:
        docx_sti = ledig_filnavn(os.path.splitext(pdf_sti)[0] + ".docx")

    status("Analyserer dokumentet ...")
    fremdrift(5)
    digital = har_tekstlag(pdf_sti)

    tmpdir = tempfile.mkdtemp(prefix="pdf2word_")
    arbeidsfil = pdf_sti
    ocr_brukt = False
    fjernet = 0

    try:
        if not digital:
            ocr_fil = os.path.join(tmpdir, "ocr.pdf")
            kjor_ocr(pdf_sti, ocr_fil, status, fremdrift, fra=10, til=55)
            arbeidsfil = ocr_fil
            ocr_brukt = True

        if fjern_marger:
            status("Fjerner topp- og bunntekst ...")
            renset = os.path.join(tmpdir, "renset.pdf")
            fjernet = fjern_topp_bunn(arbeidsfil, renset)
            if fjernet:
                arbeidsfil = renset
            fremdrift(60 if ocr_brukt else 20)

        status("Bygger Word-dokument ...")
        if ocr_brukt:
            # pdf2docx ser ikke usynlig OCR-tekst - bygg dokumentet fra tekstlaget
            bygg_docx_fra_ocr(arbeidsfil, docx_sti)
        else:
            cv = Converter(arbeidsfil)
            try:
                cv.convert(docx_sti)
            finally:
                cv.close()

        fremdrift(100)
        status("Ferdig.")
        return {"utfil": docx_sti, "ocr_brukt": ocr_brukt,
                "fjernet_blokker": fjernet}
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def kjor_i_prosess(ko, pdf_sti: str, fjern_marger: bool) -> None:
    """
    Kjøres i en egen prosess slik at minnekrasj eller heng ikke tar med seg
    programvinduet. Kommuniserer via multiprocessing-kø.
    """
    try:
        res = konverter(
            pdf_sti,
            status=lambda s: ko.put(("status", s)),
            fremdrift=lambda p: ko.put(("fremdrift", p)),
            fjern_marger=fjern_marger,
        )
        ko.put(("ferdig", res))
    except Exception as e:
        ko.put(("feil", str(e)))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Bruk: python converter.py <fil.pdf>")
        sys.exit(1)
    res = konverter(sys.argv[1], status=print, fremdrift=lambda p: None)
    print("Lagret:", res["utfil"], "| OCR:", res["ocr_brukt"])
